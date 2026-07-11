"""
DocumentExporterService (Mark X.5) — executive-grade document production.

Fully offline document generation in ORION's visual identity (crimson
#ff1a3c on deep void), with a durable export history:

    compile_docx_brief()        — an executive Word brief (python-docx; falls
                                  back to Markdown when the library is absent)
    assemble_html_deck()        — a responsive, self-contained HTML
                                  presentation (arrow keys / click to advance)
    export_report()             — a full report as Markdown + HTML (+ DOCX
                                  when available) in one export folder
    compile_presentation_deck() — assemble_html_deck persisted as an export
    get_export_history()        — everything ever exported, newest first

Design rules: no network access, no Qt imports, every filesystem write via
``asyncio.to_thread`` so the event loop never blocks, and every artefact is
recorded in ``config/exports/history.json`` and announced on the bus.
"""

from __future__ import annotations

import asyncio
import html
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .bus import OrionBus
from .constants import BASE_DIR, C, CONFIG_DIR
from .data import ToolResult
from .utils import first_line, utc_stamp

EXPORT_ROOT = BASE_DIR / "exports"
HISTORY_PATH = CONFIG_DIR / "exports" / "history.json"
HISTORY_LIMIT = 200


class DocumentExporterService:
    """Compiles briefs, reports and presentation decks entirely offline."""

    def __init__(self, bus: OrionBus, telemetry: Any | None = None) -> None:
        self.bus = bus
        self.telemetry = telemetry
        if self.telemetry is not None:
            self.telemetry.health.register("exporter")

    # ── public API ────────────────────────────────────────────────────────────

    async def compile_docx_brief(
        self,
        title: str,
        sections: list[dict[str, str]] | None = None,
        summary: str = "",
    ) -> ToolResult:
        """Executive Word brief: cover line, summary, headed sections."""
        title = self._clean_title(title)
        if not title:
            return ToolResult("A title is required for the brief.", ok=False)
        sections = self._normalise_sections(sections, summary)
        try:
            path = await asyncio.to_thread(self._write_docx_brief, title, summary, sections)
        except Exception as exc:
            return ToolResult(f"Brief compilation failed: {first_line(exc)}", ok=False)
        await self._record("docx_brief", title, path)
        return ToolResult(f"Executive brief compiled, sir: {path}")

    async def assemble_html_deck(
        self, title: str, slides: list[dict[str, str]] | None = None
    ) -> ToolResult:
        """Responsive, self-contained HTML presentation in the ORION palette."""
        title = self._clean_title(title)
        if not title:
            return ToolResult("A title is required for the deck.", ok=False)
        slides = self._normalise_slides(slides)
        try:
            path = await asyncio.to_thread(self._write_html_deck, title, slides)
        except Exception as exc:
            return ToolResult(f"Deck assembly failed: {first_line(exc)}", ok=False)
        await self._record("html_deck", title, path)
        return ToolResult(
            f"Presentation deck assembled, sir ({len(slides)} slide(s)): {path}"
        )

    async def export_report(
        self,
        title: str,
        body: str = "",
        sections: list[dict[str, str]] | None = None,
    ) -> ToolResult:
        """Full report: Markdown + styled HTML (+ DOCX when python-docx exists)."""
        title = self._clean_title(title)
        if not title:
            return ToolResult("A title is required for the report.", ok=False)
        sections = self._normalise_sections(sections, body)
        try:
            paths = await asyncio.to_thread(self._write_report, title, body, sections)
        except Exception as exc:
            return ToolResult(f"Report export failed: {first_line(exc)}", ok=False)
        await self._record("report", title, paths[0])
        return ToolResult(
            "Report exported, sir:\n" + "\n".join(f"- {p}" for p in paths)
        )

    async def compile_presentation_deck(
        self, title: str, slides: list[dict[str, str]] | None = None
    ) -> ToolResult:
        """Alias with presentation semantics — kept distinct for the tool map."""
        return await self.assemble_html_deck(title, slides)

    def get_export_history(self, limit: int = 20) -> ToolResult:
        """Everything exported, newest first."""
        entries = self._read_history()[: max(1, min(100, int(limit or 20)))]
        if not entries:
            return ToolResult("Nothing has been exported yet, sir.")
        lines = [f"Export history — {len(entries)} item(s):"]
        for e in entries:
            lines.append(f"- [{e.get('at', '')[:19]}] {e.get('kind')}: "
                         f"{e.get('title')} → {e.get('path')}")
        return ToolResult("\n".join(lines))

    # ── writers (worker threads) ──────────────────────────────────────────────

    def _export_dir(self, title: str) -> Path:
        folder = EXPORT_ROOT / f"{datetime.now():%Y-%m-%d}_{self._slug(title)}"
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    def _write_docx_brief(
        self, title: str, summary: str, sections: list[dict[str, str]]
    ) -> Path:
        folder = self._export_dir(title)
        try:
            import docx  # type: ignore
            from docx.shared import Pt, RGBColor  # type: ignore

            document = docx.Document()
            crimson = RGBColor(0xFF, 0x1A, 0x3C)
            heading = document.add_heading(title, level=0)
            for run in heading.runs:
                run.font.color.rgb = crimson
            meta = document.add_paragraph(
                f"O.R.I.O.N. executive brief — {datetime.now():%A %d %B %Y, %H:%M}"
            )
            meta.runs[0].font.size = Pt(9)
            if summary.strip():
                document.add_heading("Executive summary", level=1)
                document.add_paragraph(summary.strip())
            for section in sections:
                document.add_heading(section["heading"], level=1)
                for paragraph in section["content"].split("\n\n"):
                    if paragraph.strip():
                        document.add_paragraph(paragraph.strip())
            path = folder / f"{self._slug(title)}.docx"
            document.save(str(path))
            return path
        except ImportError:
            # Graceful offline degradation: identical structure as Markdown.
            path = folder / f"{self._slug(title)}.md"
            path.write_text(self._markdown(title, summary, sections), encoding="utf-8")
            self.bus.log.emit(
                "EXPORT: python-docx not installed; brief written as Markdown instead."
            )
            return path

    def _write_report(
        self, title: str, body: str, sections: list[dict[str, str]]
    ) -> list[Path]:
        folder = self._export_dir(title)
        slug = self._slug(title)
        paths: list[Path] = []
        markdown = self._markdown(title, body, sections)
        md_path = folder / f"{slug}.md"
        md_path.write_text(markdown, encoding="utf-8")
        paths.append(md_path)
        html_path = folder / f"{slug}.html"
        html_path.write_text(self._report_html(title, body, sections), encoding="utf-8")
        paths.append(html_path)
        try:
            paths.append(self._write_docx_brief(title, body, sections))
        except Exception:
            pass  # DOCX is a bonus; Markdown + HTML already succeeded
        return paths

    def _write_html_deck(self, title: str, slides: list[dict[str, str]]) -> Path:
        folder = self._export_dir(title)
        path = folder / f"{self._slug(title)}_deck.html"
        path.write_text(self._deck_html(title, slides), encoding="utf-8")
        return path

    # ── renderers ─────────────────────────────────────────────────────────────

    @staticmethod
    def _markdown(title: str, summary: str, sections: list[dict[str, str]]) -> str:
        lines = [f"# {title}", "",
                 f"*O.R.I.O.N. export — {datetime.now():%A %d %B %Y, %H:%M}*", ""]
        if summary.strip():
            lines += ["## Executive summary", "", summary.strip(), ""]
        for section in sections:
            lines += [f"## {section['heading']}", "", section["content"].strip(), ""]
        return "\n".join(lines)

    def _report_html(
        self, title: str, summary: str, sections: list[dict[str, str]]
    ) -> str:
        blocks = ""
        if summary.strip():
            blocks += (f"<section><h2>Executive summary</h2>"
                       f"<p>{self._para(summary)}</p></section>")
        for section in sections:
            blocks += (f"<section><h2>{html.escape(section['heading'])}</h2>"
                       f"<p>{self._para(section['content'])}</p></section>")
        return f"""<!DOCTYPE html>
<html lang="en-GB"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title><style>
:root {{ --pri: {C.PRI}; --bg: {C.BG}; --panel: {C.PANEL}; --muted: {C.MUTED}; }}
body {{ margin: 0; background: var(--bg); color: #fff;
       font-family: 'Segoe UI', system-ui, sans-serif; line-height: 1.6; }}
main {{ max-width: 860px; margin: 0 auto; padding: 48px 24px; }}
h1 {{ color: var(--pri); border-bottom: 2px solid var(--pri);
     padding-bottom: 12px; letter-spacing: 1px; }}
h2 {{ color: var(--pri); margin-top: 40px; }}
section {{ background: var(--panel); border: 1px solid {C.BORDER};
          border-radius: 8px; padding: 20px 24px; margin: 18px 0; }}
.meta {{ color: var(--muted); font-size: 0.85rem; }}
</style></head><body><main>
<h1>{html.escape(title)}</h1>
<p class="meta">O.R.I.O.N. executive report — {datetime.now():%A %d %B %Y, %H:%M}</p>
{blocks}
</main></body></html>"""

    def _deck_html(self, title: str, slides: list[dict[str, str]]) -> str:
        slide_html = ""
        for index, slide in enumerate(slides):
            slide_html += (
                f'<div class="slide{" active" if index == 0 else ""}">'
                f"<h2>{html.escape(slide['heading'])}</h2>"
                f"<div class='slide-body'>{self._bullets(slide['content'])}</div>"
                f"<div class='page'>{index + 1} / {len(slides)}</div></div>"
            )
        return f"""<!DOCTYPE html>
<html lang="en-GB"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title><style>
:root {{ --pri: {C.PRI}; --bg: {C.BG}; --panel: {C.PANEL}; --muted: {C.MUTED}; }}
* {{ box-sizing: border-box; }}
body {{ margin: 0; background: var(--bg); color: #fff; overflow: hidden;
       font-family: 'Segoe UI', system-ui, sans-serif; }}
.slide {{ display: none; position: absolute; inset: 0; padding: 6vh 8vw;
         flex-direction: column; justify-content: center; }}
.slide.active {{ display: flex; animation: fade .45s ease; }}
@keyframes fade {{ from {{ opacity: 0; transform: translateY(12px); }}
                   to {{ opacity: 1; transform: none; }} }}
h2 {{ color: var(--pri); font-size: clamp(1.6rem, 4.5vw, 3.2rem);
     border-left: 6px solid var(--pri); padding-left: 18px; }}
.slide-body {{ font-size: clamp(1rem, 2.2vw, 1.5rem); line-height: 1.7; }}
.slide-body li {{ margin: .45em 0; }}
.page {{ position: absolute; bottom: 3vh; right: 4vw; color: var(--muted); }}
.brand {{ position: fixed; top: 3vh; left: 4vw; color: var(--pri);
         letter-spacing: 3px; font-size: .8rem; }}
</style></head><body>
<div class="brand">O.R.I.O.N. — {html.escape(title.upper())}</div>
{slide_html}
<script>
const slides = document.querySelectorAll('.slide');
let current = 0;
function show(next) {{
  slides[current].classList.remove('active');
  current = (next + slides.length) % slides.length;
  slides[current].classList.add('active');
}}
document.addEventListener('keydown', e => {{
  if (['ArrowRight','PageDown',' '].includes(e.key)) show(current + 1);
  if (['ArrowLeft','PageUp'].includes(e.key)) show(current - 1);
}});
document.addEventListener('click', () => show(current + 1));
</script></body></html>"""

    # ── history ───────────────────────────────────────────────────────────────

    async def _record(self, kind: str, title: str, path: Path) -> None:
        entry = {"at": utc_stamp(), "kind": kind, "title": title, "path": str(path)}
        await asyncio.to_thread(self._append_history, entry)
        self.bus.dashboard_event.emit("export", entry)
        self.bus.log.emit(f"EXPORT: {kind} '{title}' → {path}")
        if self.telemetry is not None:
            self.telemetry.metrics.incr(f"export.{kind}")
            self.telemetry.health.beat("exporter", "OK", kind)

    def _append_history(self, entry: dict[str, Any]) -> None:
        entries = self._read_history()
        entries.insert(0, entry)
        HISTORY_PATH.parent.mkdir(parents=True, exist_ok=True)
        HISTORY_PATH.write_text(
            json.dumps(entries[:HISTORY_LIMIT], indent=2), encoding="utf-8"
        )

    def _read_history(self) -> list[dict[str, Any]]:
        try:
            data = json.loads(HISTORY_PATH.read_text(encoding="utf-8"))
            return data if isinstance(data, list) else []
        except (OSError, ValueError):
            return []

    # ── normalisation helpers ─────────────────────────────────────────────────

    @staticmethod
    def _clean_title(title: str) -> str:
        return re.sub(r"\s+", " ", str(title or "")).strip()[:120]

    @staticmethod
    def _slug(title: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")[:60] or "export"

    @staticmethod
    def _normalise_sections(
        sections: list[dict[str, str]] | None, fallback_body: str
    ) -> list[dict[str, str]]:
        out: list[dict[str, str]] = []
        for raw in sections or []:
            if not isinstance(raw, dict):
                continue
            heading = str(raw.get("heading") or raw.get("title") or "Section").strip()
            content = str(raw.get("content") or raw.get("body") or "").strip()
            if content:
                out.append({"heading": heading[:120], "content": content[:20000]})
        if not out and fallback_body.strip():
            out.append({"heading": "Overview", "content": fallback_body.strip()[:20000]})
        return out

    @staticmethod
    def _normalise_slides(
        slides: list[dict[str, str]] | None
    ) -> list[dict[str, str]]:
        out: list[dict[str, str]] = []
        for raw in slides or []:
            if not isinstance(raw, dict):
                continue
            heading = str(raw.get("heading") or raw.get("title") or "Slide").strip()
            content = str(raw.get("content") or raw.get("body") or "").strip()
            out.append({"heading": heading[:120], "content": content[:4000]})
        return out or [{"heading": "Untitled", "content": "No content supplied."}]

    @staticmethod
    def _para(text: str) -> str:
        return html.escape(text.strip()).replace("\n\n", "</p><p>").replace("\n", "<br>")

    def _bullets(self, text: str) -> str:
        lines = [l.strip("•- \t") for l in text.splitlines() if l.strip()]
        if len(lines) > 1:
            return "<ul>" + "".join(f"<li>{html.escape(l)}</li>" for l in lines) + "</ul>"
        return f"<p>{self._para(text)}</p>"
